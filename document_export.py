from io import BytesIO
import re

def create_word_document(report_content, child_name, report_id):
    """Create a Word document from the report content"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'EHC Assessment Report', 0)
        title.alignment = 1  # Center alignment
        
        # Add child info
        doc.add_heading(f'Child: {child_name}', level=2)
        doc.add_paragraph(f'Report ID: {report_id}')
        doc.add_paragraph('')  # Add spacing
        
        # Parse and add content
        lines = report_content.split('\n')
        current_paragraph = []
        
        for line in lines:
            line = line.strip()
            
            if not line:
                # Empty line - end current paragraph if exists
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                continue
            
            # Check if this looks like a heading
            if (line.isupper() or 
                line.startswith('#') or 
                re.match(r'^\d+\.', line) or
                line.endswith(':') and len(line) < 100):
                
                # Add any pending paragraph first
                if current_paragraph:
                    doc.add_paragraph(' '.join(current_paragraph))
                    current_paragraph = []
                
                # Add as heading
                heading_text = line.replace('#', '').strip()
                doc.add_heading(heading_text, level=1)
            else:
                # Regular content
                current_paragraph.append(line)
        
        # Add any remaining paragraph
        if current_paragraph:
            doc.add_paragraph(' '.join(current_paragraph))
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.getvalue()
        
    except ImportError:
        # Fallback to text if python-docx not available
        return None

def create_rich_text_document(report_content, child_name, report_id):
    """Create an RTF document as fallback"""
    rtf_header = r"""{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}
\f0\fs24 """
    
    rtf_footer = r"}"
    
    # Clean and format content
    content = report_content.replace('\n\n', r'\par\par ')
    content = content.replace('\n', r'\par ')
    
    # Add title
    rtf_content = rtf_header
    rtf_content += r"\b\fs28 EHC Assessment Report\b0\fs24\par\par "
    rtf_content += f"Child: {child_name}\\par "
    rtf_content += f"Report ID: {report_id}\\par\\par "
    rtf_content += content
    rtf_content += rtf_footer
    
    return rtf_content.encode('utf-8')